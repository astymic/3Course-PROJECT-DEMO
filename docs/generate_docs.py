"""
generate_docs.py — генерує documentation.docx та sprint2_overview.docx
Запуск: python demo/docs/generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Кольори ─────────────────────────────────────────────────────────────────
AMBER   = RGBColor(0x92, 0x40, 0x0E)   # amber-800
DARK    = RGBColor(0x1C, 0x19, 0x17)   # stone-900
GREY    = RGBColor(0x57, 0x53, 0x4E)   # stone-600
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
AMBER_L = RGBColor(0xFF, 0xFB, 0xEB)   # amber-50 bg
GREY_L  = RGBColor(0xF5, 0xF5, 0xF4)   # stone-100 bg


# ── Допоміжні функції ────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Заливка комірки таблиці."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color: str):
    """Заливка абзацу."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = AMBER
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_body(doc: Document, text: str, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)
    return p


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    set_para_bg(p, 'F5F5F4')
    return p


def add_note(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f'⚠  {text}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = AMBER
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    set_para_bg(p, 'FFFBEB')
    return p


def add_table(doc: Document, headers: list, rows: list):
    col_count = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        set_cell_bg(cell, '92400E')

    # Data rows
    for ri, row in enumerate(rows):
        bg = 'F5F5F4' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
            run.font.name = 'Calibri'
            set_cell_bg(cell, bg)

    doc.add_paragraph()  # space after table
    return t


def add_spacer(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def set_doc_margins(doc: Document):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)


# ════════════════════════════════════════════════════════════════════════════
#  ДОКУМЕНТ 1: documentation.docx
# ════════════════════════════════════════════════════════════════════════════

def build_documentation():
    doc = Document()
    set_doc_margins(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('LiLu — Взуттєва фабрика')
    tr.font.name = 'Calibri'
    tr.font.size = Pt(28)
    tr.font.bold = True
    tr.font.color.rgb = AMBER
    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('Технічна документація проєкту')
    sr.font.name = 'Calibri'
    sr.font.size = Pt(16)
    sr.font.color.rgb = GREY
    sr.font.italic = True

    doc.add_paragraph()

    # ── 1. Загальний огляд ────────────────────────────────────────────────
    add_heading(doc, '1. Загальний огляд', 1)
    add_body(doc, 'LiLu — веб-платформа для продажу взуття власного виробництва. Система включає:')
    add_bullet(doc, 'Каталог товарів з фільтрацією за категорією, кольором, розміром, сезоном')
    add_bullet(doc, 'Кошик і оформлення замовлень (Нова Пошта або кур\'єр)')
    add_bullet(doc, 'Систему авторизації через телефон або email')
    add_bullet(doc, 'Особистий кабінет покупця з історією замовлень і підтримки')
    add_bullet(doc, 'Адміністративну панель (товари, замовлення, користувачі)')
    add_bullet(doc, 'Live чат-підтримку з адмін-дашбордом у реальному часі')

    # ── 2. Структура проєкту ─────────────────────────────────────────────
    add_heading(doc, '2. Структура проєкту', 1)
    add_table(doc,
        ['Папка', 'Опис'],
        [
            ['demo/docs/', 'Документація (цей файл)'],
            ['demo/sprint1/', 'Базовий каталог і кошик'],
            ['demo/sprint2/', 'Авторизація, кабінет, чат'],
            ['demo/sprint2/app/', 'Next.js сторінки та API-маршрути'],
            ['demo/sprint2/components/', 'React-компоненти'],
            ['demo/sprint2/lib/', 'auth.ts, prisma.ts'],
            ['demo/sprint2/prisma/', 'Схема БД, міграції, seed'],
        ]
    )

    # ── 3. Технічний стек ────────────────────────────────────────────────
    add_heading(doc, '3. Технічний стек', 1)
    add_table(doc,
        ['Компонент', 'Технологія', 'Версія'],
        [
            ['Фреймворк', 'Next.js', '16.2.4'],
            ['Стилізація', 'Tailwind CSS', '3.x'],
            ['ORM', 'Prisma', '5.22'],
            ['База даних', 'SQLite', '—'],
            ['Авторизація', 'HMAC-SHA256 + httpOnly cookie', '—'],
            ['Edge Runtime', 'Web Crypto API', '—'],
            ['Реальний час', 'Polling (setInterval 3 сек)', '—'],
        ]
    )

    # ── 4. База даних ────────────────────────────────────────────────────
    add_heading(doc, '4. База даних', 1)
    add_body(doc, 'Всі моделі описані у prisma/schema.prisma. Тип сховища — SQLite.')

    add_heading(doc, 'Category', 2)
    add_table(doc, ['Поле', 'Тип', 'Опис'], [
        ['id', 'Int PK', 'Первинний ключ'],
        ['name', 'String', 'Назва категорії'],
        ['slug', 'String unique', 'URL-slug'],
        ['season', 'String', 'summer / winter / demi'],
    ])

    add_heading(doc, 'Product', 2)
    add_table(doc, ['Поле', 'Тип', 'Опис'], [
        ['id', 'Int PK', 'Первинний ключ'],
        ['name', 'String', 'Назва товару'],
        ['description', 'String', 'Опис'],
        ['price', 'Float', 'Ціна (грн)'],
        ['material', 'String', 'Матеріал'],
        ['color', 'String', 'Колір'],
        ['imageUrl', 'String', 'Шлях до зображення'],
        ['isActive', 'Boolean', 'Видимість у каталозі'],
        ['categoryId', 'Int FK', 'FK → Category'],
    ])

    add_heading(doc, 'User', 2)
    add_table(doc, ['Поле', 'Тип', 'Опис'], [
        ['id', 'Int PK', 'Первинний ключ'],
        ['name', 'String', "Ім'я"],
        ['phone', 'String?', 'Телефон (унікальний)'],
        ['email', 'String?', 'Email (унікальний)'],
        ['password', 'String', 'HMAC-SHA256 хеш'],
        ['role', 'String', 'user / admin'],
    ])

    add_heading(doc, 'Order', 2)
    add_table(doc, ['Поле', 'Тип', 'Опис'], [
        ['id', 'Int PK', 'Первинний ключ'],
        ['userId', 'Int?', 'FK → User (null для гостей)'],
        ['customerName/Phone/Email', 'String', 'Контактні дані покупця'],
        ['deliveryMethod', 'String', 'nova_poshta / courier'],
        ['npCity / npWarehouse', 'String', 'Місто і відділення НП'],
        ['paymentMethod', 'String', 'liqpay / cash_on_delivery'],
        ['paymentStatus', 'String', 'pending / paid'],
        ['status', 'String', 'new / processing / shipped / delivered / cancelled'],
        ['total', 'Float', 'Загальна сума'],
    ])

    add_heading(doc, 'ChatSession / ChatMessage', 2)
    add_table(doc, ['Поле', 'Тип', 'Опис'], [
        ['token', 'String unique', 'UUID — зберігається у клієнті'],
        ['userId', 'Int?', 'FK → User (якщо авторизований)'],
        ['status', 'String', 'open / closed'],
        ['sender', 'String', 'user / admin (у ChatMessage)'],
    ])

    # ── 5. Авторизація ───────────────────────────────────────────────────
    add_heading(doc, '5. Авторизація', 1)
    add_body(doc, 'Формат токена:')
    add_code(doc, 'base64( userId:role:expiresTimestamp:hmac_sha256_signature )')
    add_body(doc, 'Захищені маршрути:')
    add_table(doc, ['Маршрут', 'Умова доступу'], [
        ['/account/*', 'Будь-який авторизований користувач'],
        ['/admin/*', 'Тільки роль admin'],
    ])
    add_body(doc, 'Початковий адмін: admin@lilu.ua / admin123')
    add_note(doc, 'Файл proxy.ts замінює middleware.ts. Використовується Web Crypto API (Edge Runtime).')

    # ── 6. API ───────────────────────────────────────────────────────────
    add_heading(doc, '6. API Reference', 1)

    add_heading(doc, 'Авторизація', 2)
    add_table(doc, ['Метод', 'URL', 'Опис'], [
        ['POST', '/api/auth/register', 'Реєстрація'],
        ['POST', '/api/auth/login', 'Вхід'],
        ['GET', '/api/auth/me', 'Поточний користувач'],
        ['DELETE', '/api/auth/me', 'Вихід'],
        ['PUT', '/api/auth/profile', 'Оновлення профілю'],
    ])

    add_heading(doc, 'Товари', 2)
    add_table(doc, ['Метод', 'URL', 'Опис'], [
        ['GET', '/api/products', 'Список (фільтри: category, color, size, search)'],
        ['POST', '/api/products', 'Створити (admin)'],
        ['PUT', '/api/products/[id]', 'Оновити (admin)'],
        ['DELETE', '/api/products/[id]', 'Видалити (admin)'],
    ])

    add_heading(doc, 'Замовлення', 2)
    add_table(doc, ['Метод', 'URL', 'Опис'], [
        ['POST', '/api/orders', 'Створити (прив\'язує userId)'],
        ['GET', '/api/orders', 'Всі замовлення (admin)'],
        ['GET', '/api/orders/mine', 'Замовлення поточного користувача'],
        ['PUT', '/api/orders/[id]', 'Змінити статус (admin)'],
    ])

    add_heading(doc, 'Чат', 2)
    add_table(doc, ['Метод', 'URL', 'Опис'], [
        ['POST', '/api/chat/sessions', 'Створити сесію'],
        ['GET', '/api/chat/sessions', 'Свої сесії (авторизований)'],
        ['GET', '/api/chat/sessions/[id]/messages', 'Polling повідомлень'],
        ['POST', '/api/chat/sessions/[id]/messages', 'Надіслати повідомлення'],
        ['GET', '/api/admin/chat', 'Всі сесії (admin)'],
        ['POST', '/api/admin/chat/[id]', 'Відповідь адміна'],
        ['PUT', '/api/admin/chat/[id]', 'Змінити статус сесії'],
    ])

    # ── 7. Компоненти ────────────────────────────────────────────────────
    add_heading(doc, '7. Компоненти', 1)
    add_table(doc, ['Компонент', 'Опис'], [
        ['HeaderNav', 'Навігація з динамічним станом авторизації (Гість / User / Admin)'],
        ['CartDropdown', 'Плаваючий кошик з лічильником'],
        ['ChatWidget', 'Floating чат-підтримка: FAQ бот + живий спеціаліст'],
        ['FilterPanel', 'Панель фільтрів каталогу'],
        ['ProductCard', 'Картка товару з вибором розміру'],
        ['ProductTable', 'Таблиця товарів для адмін-панелі'],
    ])

    # ── 8. Налаштування ──────────────────────────────────────────────────
    add_heading(doc, '8. Налаштування та запуск', 1)
    add_heading(doc, 'Змінні оточення (.env)', 2)
    add_code(doc, 'DATABASE_URL="file:./prisma/dev.db"')
    add_code(doc, 'AUTH_SECRET="your-secret-key-change-in-production"')
    add_code(doc, 'NOVA_POSHTA_API_KEY="your-np-api-key"')
    add_heading(doc, 'Перший запуск', 2)
    add_code(doc, 'cd demo/sprint2')
    add_code(doc, 'npm install')
    add_code(doc, 'npx prisma migrate dev')
    add_code(doc, 'node prisma/seed.js')
    add_code(doc, 'npm run dev')

    # ── 9. Ролі ──────────────────────────────────────────────────────────
    add_heading(doc, '9. Ролі та доступ', 1)
    add_table(doc, ['Роль', 'Можливості'], [
        ['Гість', 'Каталог, кошик, оформлення замовлення, чат (анонімно)'],
        ['user', 'Все вище + кабінет, замовлення, профіль, збережений чат'],
        ['admin', 'Все вище + адмін-панель: товари, замовлення, користувачі, чат'],
    ])

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Документація оновлюється після кожного спринту  •  LiLu E-Commerce  •  2026')
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.font.italic = True

    out = os.path.join(DOCS_DIR, 'documentation.docx')
    doc.save(out)
    print(f'OK  {out}')


# ════════════════════════════════════════════════════════════════════════════
#  ДОКУМЕНТ 2: sprint2_overview.docx
# ════════════════════════════════════════════════════════════════════════════

def build_sprint2_overview():
    doc = Document()
    set_doc_margins(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('LiLu E-Commerce — Sprint 2')
    tr.font.name = 'Calibri'; tr.font.size = Pt(26); tr.font.bold = True; tr.font.color.rgb = AMBER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('Огляд виконаної роботи  •  квітень 2026 р.')
    sr.font.name = 'Calibri'; sr.font.size = Pt(13); sr.font.color.rgb = GREY; sr.font.italic = True

    doc.add_paragraph()
    add_body(doc, 'Sprint 2 — впровадження повноцінної системи авторизації, особистого кабінету, адміністративних інструментів і live чат-підтримки.')

    # ── Для покупця ──────────────────────────────────────────────────────
    add_heading(doc, 'Нові можливості для покупця', 1)

    add_heading(doc, 'Реєстрація і вхід', 2)
    add_bullet(doc, 'Реєстрація через телефон або email (або одночасно обидва)')
    add_bullet(doc, 'Вхід за телефоном або email + пароль')
    add_bullet(doc, 'Сесія зберігається 7 днів (httpOnly cookie)')

    add_heading(doc, 'Особистий кабінет (/account)', 2)
    add_table(doc, ['Вкладка', 'Що показує'], [
        ['Замовлення', 'Повна історія з деталями, статусами, методами доставки'],
        ['Профіль', "Редагування імені, телефону, email; зміна пароля"],
        ['Підтримка', 'Архів звернень до чату з повними діалогами'],
    ])

    add_heading(doc, 'Розумна рекламація замовлень', 2)
    add_body(doc, 'Якщо покупець оформив замовлення як гість, а потім зареєструвався з тим самим телефоном або email — всі попередні замовлення автоматично прив\'язуються до його акаунту. На сторінці кабінету відображається банер з кількістю знайдених замовлень.')

    add_heading(doc, 'Автозаповнення в кошику', 2)
    add_body(doc, 'Кнопка «Заповнити з профілю» на сторінці оформлення замовлення — підтягує ім\'я, телефон та email авторизованого користувача.')

    add_heading(doc, 'Чат-підтримка', 2)
    add_bullet(doc, 'Floating кнопка 💬 на всіх сторінках сайту')
    add_bullet(doc, '6 швидких питань (FAQ chips) з автоматичними відповідями бота')
    add_bullet(doc, 'Кнопка «Поговорити зі спеціалістом» — підключає живого оператора')
    add_bullet(doc, 'При підключенні спеціаліста — увесь FAQ-контекст автоматично надсилається адміну')
    add_bullet(doc, 'Гостьовий чат очищується при закритті вкладки')
    add_bullet(doc, 'Чат авторизованого — зберігається між сесіями')
    add_bullet(doc, 'Кнопка «Почати новий чат» при закритому зверненні')

    # ── Для адміністратора ───────────────────────────────────────────────
    add_heading(doc, 'Нові можливості для адміністратора', 1)

    add_heading(doc, 'Управління користувачами (/admin/users)', 2)
    add_bullet(doc, 'Список всіх зареєстрованих користувачів')
    add_bullet(doc, 'Призначення ролі admin через інтерфейс')

    add_heading(doc, 'Чат-дашборд (/admin/chat)', 2)
    add_bullet(doc, 'Sidebar зі списком звернень (фільтр: Відкриті / Закриті / Всі)')
    add_bullet(doc, 'Повна переписка включно з FAQ-контекстом до підключення')
    add_bullet(doc, 'Відповіді в реальному часі — оновлення кожні 3 секунди')
    add_bullet(doc, 'Закриття / відновлення теми спілкування')

    # ── Технічні рішення ─────────────────────────────────────────────────
    add_heading(doc, 'Ключові технічні рішення', 1)
    add_table(doc, ['Задача', 'Рішення'], [
        ['Edge Runtime несумісний з Node.js crypto', 'Перехід на Web Crypto API (crypto.subtle)'],
        ['middleware.ts застарів у Next.js 16', 'Замінено на proxy.ts'],
        ['НП показувала лише 50 відділень', 'Ліміт збільшено до 500'],
        ['Дублювання повідомлень у чаті', 'seenIds Set — кожен ID відстежується та не додається двічі'],
        ['Гостьові замовлення недоступні після реєстрації', 'Автоматична рекламація при реєстрації (phone або email match)'],
    ])

    # ── Що не змінювалося ─────────────────────────────────────────────────
    add_heading(doc, 'Що не змінювалося', 1)
    add_bullet(doc, 'Каталог, фільтри, кошик — без змін (Sprint 1)')
    add_bullet(doc, 'LiqPay — залишається заглушкою (без реальної інтеграції)')
    add_bullet(doc, 'Підтвердження телефону/email — не реалізовано (навчальний проєкт)')

    # ── Початковий адмін ──────────────────────────────────────────────────
    add_heading(doc, 'Початковий адміністратор', 1)
    add_table(doc, ['Поле', 'Значення'], [
        ['Email', 'admin@lilu.ua'],
        ['Пароль', 'admin123'],
        ['Роль', 'admin'],
    ])
    add_body(doc, 'Команда для створення: node prisma/seed.js')

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LiLu E-Commerce  •  Sprint 2  •  2026')
    r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.italic = True

    out = os.path.join(DOCS_DIR, 'sprint2_overview.docx')
    doc.save(out)
    print(f'OK  {out}')


# ── Точка входу ──────────────────────────────────────────────────────────────
if __name__ == '__main__':
    build_documentation()
    build_sprint2_overview()
    print('\nВсе готово! Файли збережено у demo/docs/')
